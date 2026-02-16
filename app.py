import streamlit as st
import requests
import json
import re
import os
import io
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from youtube_transcript_api import YouTubeTranscriptApi
from dotenv import load_dotenv


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
load_dotenv()

def get_config():
    """Get API configuration from environment or Streamlit secrets."""
    api_key = os.getenv("EURI_API_KEY") or st.secrets.get("EURI_API_KEY", "")
    api_url = os.getenv("EURI_URL") or st.secrets.get(
        "EURI_URL", "https://api.euron.one/api/v1/euri/chat/completions"
    )
    return api_key, api_url

# ---------------------------------------------------------------------------
# Page Config
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="VideoInsight AI — Video Summarizer",
    page_icon="🎬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------------------------------------------------------------------------
# Custom CSS
# ---------------------------------------------------------------------------
st.markdown("""
<style>
/* ---------- Google Font ---------- */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

html, body, [class*="st-"] {
    font-family: 'Inter', sans-serif;
}

/* ---------- Glassmorphism cards ---------- */
.glass-card {
    background: rgba(255,255,255,0.04);
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 16px;
    padding: 24px;
    margin-bottom: 20px;
    backdrop-filter: blur(12px);
    -webkit-backdrop-filter: blur(12px);
    transition: transform 0.25s ease, box-shadow 0.25s ease;
}
.glass-card:hover {
    transform: translateY(-2px);
    box-shadow: 0 8px 32px rgba(124,58,237,0.18);
}

/* ---------- Hero header ---------- */
.hero-header {
    text-align: center;
    padding: 40px 20px 20px;
}
.hero-header h1 {
    font-size: 2.8rem;
    font-weight: 800;
    background: linear-gradient(135deg, #7C3AED 0%, #EC4899 50%, #F59E0B 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    margin-bottom: 8px;
}
.hero-header p {
    color: #94A3B8;
    font-size: 1.1rem;
    font-weight: 300;
}

/* ---------- Feature badges ---------- */
.feature-badges {
    display: flex;
    gap: 12px;
    justify-content: center;
    flex-wrap: wrap;
    margin: 16px 0 30px;
}
.badge {
    background: linear-gradient(135deg, rgba(124,58,237,0.15), rgba(236,72,153,0.10));
    border: 1px solid rgba(124,58,237,0.3);
    border-radius: 24px;
    padding: 6px 18px;
    font-size: 0.82rem;
    color: #C4B5FD;
    font-weight: 500;
    letter-spacing: 0.02em;
}

/* ---------- Section titles ---------- */
.section-title {
    font-size: 1.3rem;
    font-weight: 700;
    color: #E2E8F0;
    margin-bottom: 16px;
    display: flex;
    align-items: center;
    gap: 10px;
}

/* ---------- Stat cards ---------- */
.stat-row {
    display: flex;
    gap: 14px;
    margin-bottom: 20px;
}
.stat-card {
    flex: 1;
    background: rgba(124,58,237,0.08);
    border: 1px solid rgba(124,58,237,0.18);
    border-radius: 12px;
    padding: 16px;
    text-align: center;
}
.stat-value {
    font-size: 1.5rem;
    font-weight: 700;
    color: #A78BFA;
}
.stat-label {
    font-size: 0.78rem;
    color: #94A3B8;
    text-transform: uppercase;
    letter-spacing: 0.06em;
    margin-top: 2px;
}

/* ---------- Summary content ---------- */
.summary-content {
    background: rgba(255,255,255,0.02);
    border: 1px solid rgba(255,255,255,0.06);
    border-radius: 12px;
    padding: 24px;
    line-height: 1.6;
    color: #CBD5E1;
    font-size: 0.95rem;
}
.summary-content h1, .summary-content h2, .summary-content h3 {
    color: #E2E8F0;
    margin-top: 16px;
    margin-bottom: 8px;
}
.summary-content h2 {
    font-size: 1.15rem;
    border-bottom: 1px solid rgba(124,58,237,0.2);
    padding-bottom: 6px;
}
.summary-content ul { padding-left: 20px; }
.summary-content li { margin-bottom: 4px; }
.summary-content p { margin-bottom: 10px; }

/* ---------- Chat messages ---------- */
.chat-user {
    background: rgba(124,58,237,0.12);
    border: 1px solid rgba(124,58,237,0.25);
    border-radius: 12px;
    padding: 14px 18px;
    margin-bottom: 12px;
    color: #DDD6FE;
}
.chat-ai {
    background: rgba(236,72,153,0.06);
    border: 1px solid rgba(236,72,153,0.15);
    border-radius: 12px;
    padding: 14px 18px;
    margin-bottom: 12px;
    color: #E2E8F0;
}

/* ---------- URL input styling ---------- */
.stTextInput > div > div > input {
    background: rgba(255,255,255,0.04) !important;
    border: 1px solid rgba(124,58,237,0.3) !important;
    border-radius: 12px !important;
    color: #E2E8F0 !important;
    padding: 14px 18px !important;
    font-size: 1rem !important;
    transition: border-color 0.3s ease !important;
}
.stTextInput > div > div > input:focus {
    border-color: #7C3AED !important;
    box-shadow: 0 0 0 2px rgba(124,58,237,0.2) !important;
}

/* ---------- Button styling ---------- */
.stButton > button {
    border-radius: 12px !important;
    font-weight: 600 !important;
    padding: 10px 28px !important;
    transition: all 0.3s ease !important;
}

/* ---------- Download button ---------- */
.stDownloadButton > button {
    background: linear-gradient(135deg, #7C3AED, #EC4899) !important;
    color: white !important;
    border: none !important;
    border-radius: 12px !important;
    font-weight: 600 !important;
    padding: 10px 28px !important;
}

/* ---------- Sidebar ---------- */
section[data-testid="stSidebar"] {
    background: rgba(15,15,26,0.95);
    border-right: 1px solid rgba(255,255,255,0.05);
}

/* ---------- Tabs ---------- */
.stTabs [data-baseweb="tab-list"] {
    gap: 6px;
}
.stTabs [data-baseweb="tab"] {
    border-radius: 10px;
    padding: 8px 20px;
    font-weight: 500;
}

/* ---------- Expander ---------- */
.streamlit-expanderHeader {
    font-weight: 600;
    color: #C4B5FD;
}

/* ---------- Footer ---------- */
.footer {
    text-align: center;
    padding: 30px 0 10px;
    color: #475569;
    font-size: 0.8rem;
}
</style>
""", unsafe_allow_html=True)


# ---------------------------------------------------------------------------
# Utility — Video ID extraction
# ---------------------------------------------------------------------------
def extract_video_id(url: str) -> str | None:
    """Extract YouTube video ID from various URL formats."""
    patterns = [
        r'(?:youtube\.com/watch\?v=|youtu\.be/|youtube\.com/embed/|youtube\.com/v/|youtube\.com/shorts/)([a-zA-Z0-9_-]{11})',
        r'(?:youtube\.com/watch\?.*v=)([a-zA-Z0-9_-]{11})',
    ]
    for pat in patterns:
        match = re.search(pat, url)
        if match:
            return match.group(1)
    return None


def get_video_info(video_id: str) -> dict:
    """Fetch basic video metadata via oembed (no API key needed)."""
    try:
        resp = requests.get(
            f"https://www.youtube.com/oembed?url=https://www.youtube.com/watch?v={video_id}&format=json",
            timeout=10,
        )
        if resp.status_code == 200:
            data = resp.json()
            return {
                "title": data.get("title", "Unknown Title"),
                "author": data.get("author_name", "Unknown"),
                "thumbnail": f"https://img.youtube.com/vi/{video_id}/maxresdefault.jpg",
            }
    except Exception:
        pass
    return {
        "title": "Video",
        "author": "Unknown",
        "thumbnail": f"https://img.youtube.com/vi/{video_id}/hqdefault.jpg",
    }


# ---------------------------------------------------------------------------
# Utility — Transcript fetching (multi-method with fallbacks)
# ---------------------------------------------------------------------------
def _fetch_via_ytdlp(video_id: str) -> str:
    """Method 1: Use yt-dlp to grab subtitles — most reliable, bypasses blocks."""
    import yt_dlp
    import tempfile
    import glob
    import json

    with tempfile.TemporaryDirectory() as tmpdir:
        url = f"https://www.youtube.com/watch?v={video_id}"
        ydl_opts = {
            'skip_download': True,
            'writesubtitles': True,
            'writeautomaticsub': True,
            'subtitleslangs': ['en', 'hi', 'en-orig', 'a.en'],
            'subtitlesformat': 'json3',
            'outtmpl': os.path.join(tmpdir, '%(id)s'),
            'quiet': True,
            'no_warnings': True,
        }
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])

        # Find any downloaded subtitle file
        sub_files = glob.glob(os.path.join(tmpdir, "*.json3"))
        if not sub_files:
            # Also check for vtt files
            sub_files = glob.glob(os.path.join(tmpdir, "*.vtt"))

        if not sub_files:
            return ""

        # Parse json3 subtitle format
        sub_file = sub_files[0]
        if sub_file.endswith('.json3'):
            with open(sub_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            lines = []
            for event in data.get('events', []):
                segs = event.get('segs', [])
                text = ''.join(s.get('utf8', '') for s in segs).strip()
                if text and text != '\n':
                    lines.append(text)
            return ' '.join(lines)
        else:
            # Parse VTT
            with open(sub_file, 'r', encoding='utf-8') as f:
                content = f.read()
            import re as _re
            # Remove VTT headers and timestamps
            lines = []
            for line in content.split('\n'):
                line = line.strip()
                if not line or line.startswith('WEBVTT') or '-->' in line or _re.match(r'^\d+$', line):
                    continue
                # Remove HTML tags
                line = _re.sub(r'<[^>]+>', '', line)
                if line:
                    lines.append(line)
            return ' '.join(lines)


def _fetch_via_transcript_api(video_id: str) -> str:
    """Method 2: Use youtube-transcript-api — may be blocked by YouTube."""
    try:
        ytt = YouTubeTranscriptApi()
        result = ytt.fetch(video_id)
        lines = [snippet.text for snippet in result]
        if not lines:
            return ""
        return " ".join(lines)
    except Exception:
        return ""


def _fetch_via_whisper(video_id: str) -> str:
    """Method 3: Download audio with yt-dlp and transcribe locally with Whisper."""
    import yt_dlp
    import tempfile

    with tempfile.TemporaryDirectory() as tmpdir:
        url = f"https://www.youtube.com/watch?v={video_id}"
        audio_path = os.path.join(tmpdir, "audio.mp3")
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': os.path.join(tmpdir, 'audio.%(ext)s'),
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '128',
            }],
            'quiet': True,
            'no_warnings': True,
        }
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])

        # Find the downloaded audio file
        import glob
        audio_files = glob.glob(os.path.join(tmpdir, "audio.*"))
        if not audio_files:
            return ""

        audio_file = audio_files[0]

        # Transcribe with faster-whisper
        try:
            from faster_whisper import WhisperModel
            model = WhisperModel("base", device="cpu", compute_type="int8")
            segments, _ = model.transcribe(audio_file, beam_size=1)
            return " ".join(seg.text.strip() for seg in segments)
        except ImportError:
            pass

        # Fallback to openai-whisper
        try:
            import whisper
            model = whisper.load_model("base")
            result = model.transcribe(audio_file, fp16=False)
            return result.get("text", "")
        except ImportError:
            pass

        return ""


def fetch_transcript(video_id: str) -> str:
    """Fetch transcript using multiple methods with automatic fallback.

    Order: yt-dlp subtitles → youtube-transcript-api → Whisper local transcription.
    """
    errors = []

    # Method 1: yt-dlp subtitles (most reliable, no API blocks)
    try:
        st.write("📡 Trying subtitle extraction (Method 1)...")
        text = _fetch_via_ytdlp(video_id)
        if text and len(text.strip()) > 20:
            st.write("✅ Got transcript via subtitle extraction!")
            return text
    except Exception as e:
        errors.append(f"yt-dlp subtitles: {e}")

    # Method 2: youtube-transcript-api (may be blocked)
    try:
        st.write("📡 Trying transcript API (Method 2)...")
        text = _fetch_via_transcript_api(video_id)
        if text and len(text.strip()) > 20:
            st.write("✅ Got transcript via API!")
            return text
    except Exception as e:
        errors.append(f"transcript API: {e}")

    # Method 3: Download audio + Whisper (slowest but always works)
    try:
        st.write("🎙️ Downloading audio for local transcription (Method 3 — this may take a minute)...")
        text = _fetch_via_whisper(video_id)
        if text and len(text.strip()) > 20:
            st.write("✅ Got transcript via local AI transcription!")
            return text
    except Exception as e:
        errors.append(f"Whisper: {e}")

    raise RuntimeError(
        "Could not fetch transcript using any method. "
        f"Details: {'; '.join(errors) if errors else 'No subtitles/captions available.'}"
    )


# ---------------------------------------------------------------------------
# Utility — Transcribe local video/audio file
# ---------------------------------------------------------------------------
def transcribe_local_file(uploaded_file, model_size: str = "base") -> str:
    """Extract audio from a local video/audio file and transcribe it.

    Uses OpenAI Whisper (via faster-whisper / CTranslate2) for fast,
    accurate, fully-local transcription — no API calls needed.
    """
    from faster_whisper import WhisperModel, BatchedInferencePipeline

    suffix = os.path.splitext(uploaded_file.name)[1].lower()
    audio_extensions = [".wav", ".mp3", ".flac", ".ogg", ".m4a", ".aac"]
    video_extensions = [".mp4", ".mkv", ".avi", ".mov", ".webm", ".wmv", ".flv", ".mpeg4"]

    all_supported = video_extensions + audio_extensions
    if suffix not in all_supported:
        raise RuntimeError(
            f"Unsupported file format '{suffix}'. "
            f"Supported: {', '.join(all_supported)}"
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        # Save uploaded file to disk (faster-whisper reads from path)
        input_path = os.path.join(tmpdir, f"input{suffix}")
        with open(input_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # Auto-detect GPU for maximum speed
        device = "cpu"
        compute_type = "int8"
        try:
            import ctranslate2
            if "cuda" in ctranslate2.get_supported_compute_types("cuda"):
                device = "cuda"
                compute_type = "float16"
                st.write("🚀 **NVIDIA GPU detected** — using CUDA for fast transcription!")
        except Exception:
            pass

        if device == "cpu":
            st.write("💻 Running on CPU — transcription may take a few minutes for long videos")

        # Load Whisper model (cached after first download)
        st.write(f"🧠 Loading Whisper AI model ({model_size})...")
        model = WhisperModel(model_size, device=device, compute_type=compute_type)

        # Use Batched Inference for parallel processing
        # This splits audio into chunks and processes them in parallel across CPU cores
        batched_model = BatchedInferencePipeline(model=model)

        st.write(f"🎙️ Running parallel transcription (Batch Size: 16) on {device.upper()}...")
        
        segments, info = batched_model.transcribe(
            input_path,
            batch_size=16,
            beam_size=1,
            condition_on_previous_text=False,
        )

        duration_min = int(info.duration // 60)
        duration_sec = int(info.duration % 60)
        st.write(f"⏱️ Audio duration: **{duration_min}m {duration_sec}s** — Language: **{info.language}** (confidence {info.language_probability:.0%})")

        progress_bar = st.progress(0, text="Transcribing...")
        transcripts: list[str] = []
        for segment in segments:
            transcripts.append(segment.text.strip())
            if info.duration > 0:
                progress_bar.progress(
                    min(segment.end / info.duration, 1.0),
                    text=f"Transcribed {int(segment.end)}s / {int(info.duration)}s",
                )
        progress_bar.progress(1.0, text="✅ Transcription complete!")

        if not transcripts:
            raise RuntimeError(
                "Could not understand the audio. The speech may be unclear "
                "or the file may have no speech."
            )
        return " ".join(transcripts)


# ---------------------------------------------------------------------------
# Utility — EURI API call
# ---------------------------------------------------------------------------
def call_euri_api(messages: list[dict], api_key: str, api_url: str) -> str:
    """Call the EURI chat completions API."""
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    payload = {
        "model": "gpt-4.1-nano",
        "messages": messages,
        "temperature": 0.5,
        "max_tokens": 4096,
    }
    try:
        resp = requests.post(api_url, headers=headers, json=payload, timeout=120)
        resp.raise_for_status()
        data = resp.json()
        return data["choices"][0]["message"]["content"]
    except requests.exceptions.HTTPError as e:
        raise RuntimeError(f"API error ({resp.status_code}): {resp.text[:300]}")
    except Exception as e:
        raise RuntimeError(f"Failed to call EURI API: {e}")


# ---------------------------------------------------------------------------
# Utility — Summarisation prompts
# ---------------------------------------------------------------------------
SUMMARY_PROMPTS = {
    "Brief": (
        "You are an expert content summarizer. Provide a BRIEF summary of the following video transcript. "
        "Include:\n"
        "1. **Overview** — 2-3 sentences about what the video covers\n"
        "2. **Key Takeaways** — 3-5 bullet points of the most important points\n"
        "3. **Conclusion** — 1-2 sentences wrapping up\n\n"
        "Keep it concise and to the point. Use clear language.\n\n"
        "TRANSCRIPT:\n{transcript}"
    ),
    "Standard": (
        "You are an expert content summarizer. Provide a DETAILED summary of the following video transcript. "
        "Include:\n"
        "1. **Overview** — A paragraph describing what the video covers, who it's for, and why it matters\n"
        "2. **Key Takeaways** — 5-8 important bullet points\n"
        "3. **Detailed Notes** — Organized section-by-section breakdown of the content\n"
        "4. **Conclusion** — Summary of key learnings and recommended next steps\n\n"
        "Use markdown formatting. Be thorough but clear.\n\n"
        "TRANSCRIPT:\n{transcript}"
    ),
    "Detailed": (
        "You are an expert content summarizer and note-taker. Provide an EXHAUSTIVE and COMPREHENSIVE summary "
        "of the following video transcript. Include:\n"
        "1. **Overview** — Detailed paragraph on the video topic, target audience, and context\n"
        "2. **Key Takeaways** — 8-12 important bullet points\n"
        "3. **Detailed Notes** — Full section-by-section breakdown with sub-points, examples, and quotes\n"
        "4. **Important Quotes** — Notable direct quotes from the speaker\n"
        "5. **Action Items** — Specific actionable steps or recommendations mentioned\n"
        "6. **Conclusion** — Comprehensive wrap-up with key learnings\n\n"
        "Use rich markdown formatting. Be as thorough as possible.\n\n"
        "TRANSCRIPT:\n{transcript}"
    ),
}


# ---------------------------------------------------------------------------
# Utility — Word document generation
# ---------------------------------------------------------------------------
def create_word_doc(title: str, author: str, summary: str) -> io.BytesIO:
    """Generate a Word document from the summary."""
    doc = Document()

    # -- Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(124, 58, 237)
    run.font.bold = True

    # -- Metadata
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(f"Channel: {author}  •  Generated: {datetime.now().strftime('%B %d, %Y %I:%M %p')}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(148, 163, 184)

    doc.add_paragraph("")  # spacer

    # -- Summary content  (parse markdown-ish text)
    for line in summary.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            for r in h.runs:
                r.font.color.rgb = RGBColor(124, 58, 237)
        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            for r in h.runs:
                r.font.color.rgb = RGBColor(124, 58, 237)
        elif stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            for r in h.runs:
                r.font.color.rgb = RGBColor(124, 58, 237)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            # Handle bold fragments
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    p.add_run(part)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    p.add_run(part)

    # -- Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Generated by VideoInsight AI")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(100, 116, 139)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ===================================================================
# MAIN APP
# ===================================================================
def main():
    api_key, api_url = get_config()

    # ---- Hero Header ----
    st.markdown("""
    <div class="hero-header">
        <h1>🎬 VideoInsight AI</h1>
        <p>Transform any YouTube video into a structured, downloadable summary in seconds</p>
    </div>
    <div class="feature-badges">
        <span class="badge">🔍 Smart Transcript Extraction</span>
        <span class="badge">🤖 AI-Powered Summaries</span>
        <span class="badge">📥 Word Document Export</span>
        <span class="badge">💬 Chat with Video</span>
        <span class="badge">⚡ Free & Open Source</span>
    </div>
    """, unsafe_allow_html=True)

    # ---- Sidebar ----
    with st.sidebar:
        st.markdown("### ⚙️ Settings")
        
        # Model Selection
        model_selection = st.selectbox(
            "Transcription Model",
            ["Fast (Tiny)", "Balanced (Base)", "Accurate (Small)"],
            index=1,
            help="Tiny: Fastest • Base: Good balance • Small: Most accurate but slower",
        )
        model_map = {"Fast (Tiny)": "tiny", "Balanced (Base)": "base", "Accurate (Small)": "small"}
        selected_model = model_map[model_selection]

        summary_length = st.selectbox(
            "Summary Length",
            ["Brief", "Standard", "Detailed"],
            index=1,
            help="Brief: Quick overview • Standard: Balanced • Detailed: Comprehensive notes",
        )
        st.markdown("---")
        st.markdown("### 📖 How to Use")
        st.markdown(
            "1. Paste a **YouTube** URL or **upload** a video/audio file\n"
            "2. Click **Summarize**\n"
            "3. Read the summary or **download** as Word\n"
            "4. Use **Chat** to ask questions\n"
        )
        st.markdown("---")
        st.markdown(
            "<div style='text-align:center;color:#64748B;font-size:0.75rem;'>"
            "Built with ❤️ using Streamlit<br>Powered by EURI AI</div>",
            unsafe_allow_html=True,
        )

    # ---- Input Mode Tabs ----
    input_tab_url, input_tab_upload = st.tabs(["🔗 YouTube URL", "📁 Upload Video/Audio"])

    video_url = ""
    uploaded_file = None

    with input_tab_url:
        col1, col2 = st.columns([5, 1])
        with col1:
            video_url = st.text_input(
                "🔗 Paste YouTube Video URL",
                placeholder="https://www.youtube.com/watch?v=...",
                label_visibility="collapsed",
            )
        with col2:
            summarize_url_clicked = st.button("🚀 Summarize", use_container_width=True, type="primary", key="btn_url")

    with input_tab_upload:
        uploaded_file = st.file_uploader(
            "Upload a video or audio file",
            type=["mp4", "mkv", "avi", "mov", "webm", "wmv", "flv", "wav", "mp3", "flac", "ogg", "m4a", "aac"],
            help="Supported: MP4, MKV, AVI, MOV, WebM, WAV, MP3, FLAC, OGG, M4A (max ~200MB)",
        )
        summarize_upload_clicked = st.button("🚀 Summarize File", use_container_width=True, type="primary", key="btn_upload")

    summarize_clicked = summarize_url_clicked if video_url else False
    summarize_file_clicked = summarize_upload_clicked if uploaded_file else False

    # ---- Session State Init ----
    if "summary" not in st.session_state:
        st.session_state.summary = ""
    if "transcript" not in st.session_state:
        st.session_state.transcript = ""
    if "video_info" not in st.session_state:
        st.session_state.video_info = {}
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    # ---- Process YouTube URL ----
    if summarize_clicked and video_url:
        if not api_key:
            st.error("⚠️ EURI API key not configured. Add `EURI_API_KEY` to your `.env` file or Streamlit secrets.")
            st.stop()

        video_id = extract_video_id(video_url)
        if not video_id:
            st.error("❌ Invalid YouTube URL. Please paste a valid link (e.g. https://www.youtube.com/watch?v=...).")
            st.stop()

        # Fetch info + transcript
        with st.status("🔄 Processing video...", expanded=True) as status:
            st.write("📡 Fetching video information...")
            info = get_video_info(video_id)
            st.session_state.video_info = info

            st.write("📝 Extracting transcript...")
            try:
                transcript = fetch_transcript(video_id)
                if not transcript.strip():
                    st.error("😕 No transcript/captions found for this video. Only videos with captions can be summarized.")
                    st.stop()
                st.session_state.transcript = transcript
            except RuntimeError as e:
                st.error(f"❌ {e}")
                st.stop()

            st.write(f"🤖 Generating **{summary_length}** summary with AI...")
            prompt = SUMMARY_PROMPTS[summary_length].format(
                transcript=transcript[:12000]  # limit to avoid token overflow
            )
            try:
                summary = call_euri_api(
                    [{"role": "user", "content": prompt}],
                    api_key,
                    api_url,
                )
                st.session_state.summary = summary
                st.session_state.chat_history = []  # reset chat for new video
            except RuntimeError as e:
                st.error(f"❌ AI summarization failed: {e}")
                st.stop()

            status.update(label="✅ Summary ready!", state="complete", expanded=False)

    # ---- Process Uploaded File ----
    if summarize_file_clicked and uploaded_file:
        if not api_key:
            st.error("⚠️ EURI API key not configured. Add `EURI_API_KEY` to your `.env` file or Streamlit secrets.")
            st.stop()

        with st.status("🔄 Processing your file...", expanded=True) as status:
            st.write(f"📁 Reading file: **{uploaded_file.name}**")
            st.session_state.video_info = {
                "title": os.path.splitext(uploaded_file.name)[0],
                "author": "Local File",
                "thumbnail": "",
            }

            st.write("🎙️ Transcribing audio (this may take a moment)...")
            try:
                transcript = transcribe_local_file(uploaded_file, model_size=selected_model)
                if not transcript.strip():
                    st.error("😕 Could not extract any speech from this file.")
                    st.stop()
                st.session_state.transcript = transcript
            except RuntimeError as e:
                st.error(f"❌ {e}")
                st.stop()

            st.write(f"🤖 Generating **{summary_length}** summary with AI...")
            prompt = SUMMARY_PROMPTS[summary_length].format(
                transcript=transcript[:12000]
            )
            try:
                summary = call_euri_api(
                    [{"role": "user", "content": prompt}],
                    api_key,
                    api_url,
                )
                st.session_state.summary = summary
                st.session_state.chat_history = []
            except RuntimeError as e:
                st.error(f"❌ AI summarization failed: {e}")
                st.stop()

            status.update(label="✅ Summary ready!", state="complete", expanded=False)

    # ---- Display Results ----
    if st.session_state.summary:
        info = st.session_state.video_info
        summary = st.session_state.summary
        transcript = st.session_state.transcript

        # Video info card
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        col_thumb, col_info = st.columns([1, 2])
        with col_thumb:
            thumbnail = info.get("thumbnail", "")
            if thumbnail:
                st.image(thumbnail, use_container_width=True)
            else:
                st.markdown(
                    '<div style="background:rgba(124,58,237,0.1);border-radius:12px;'
                    'padding:40px;text-align:center;font-size:3rem;">📁</div>',
                    unsafe_allow_html=True,
                )
        with col_info:
            st.markdown(f"### 🎥 {info.get('title', 'Video')}")
            st.markdown(f"**Channel:** {info.get('author', 'Unknown')}")

            word_count = len(transcript.split())
            summary_words = len(summary.split())
            est_duration = max(1, word_count // 150)

            st.markdown(f"""
            <div class="stat-row">
                <div class="stat-card">
                    <div class="stat-value">{word_count:,}</div>
                    <div class="stat-label">Transcript Words</div>
                </div>
                <div class="stat-card">
                    <div class="stat-value">{summary_words:,}</div>
                    <div class="stat-label">Summary Words</div>
                </div>
                <div class="stat-card">
                    <div class="stat-value">~{est_duration} min</div>
                    <div class="stat-label">Est. Video Length</div>
                </div>
                <div class="stat-card">
                    <div class="stat-value">{max(1, round(100 - (summary_words/max(word_count,1))*100))}%</div>
                    <div class="stat-label">Compression</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        # Tabs: Summary | Transcript | Chat
        tab_summary, tab_transcript, tab_chat = st.tabs(["📋 Summary", "📄 Full Transcript", "💬 Chat with Video"])

        with tab_summary:
            st.markdown(f'<div class="summary-content">{_md_to_html(summary)}</div>', unsafe_allow_html=True)

            st.markdown("")

            # -- Action buttons
            btn_col1, btn_col2, btn_col3 = st.columns(3)

            with btn_col1:
                word_buf = create_word_doc(
                    info.get("title", "Video Summary"),
                    info.get("author", "Unknown"),
                    summary,
                )
                safe_title = re.sub(r'[^\w\s-]', '', info.get("title", "summary"))[:50].strip()
                st.download_button(
                    label="📥 Download Word (.docx)",
                    data=word_buf,
                    file_name=f"{safe_title}_summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

            with btn_col2:
                st.download_button(
                    label="📄 Download Text (.txt)",
                    data=summary,
                    file_name=f"{safe_title}_summary.txt",
                    mime="text/plain",
                    use_container_width=True,
                )

            with btn_col3:
                if st.button("📋 Copy Summary", use_container_width=True):
                    st.code(summary, language=None)
                    st.success("Summary displayed above — select & copy!")

        with tab_transcript:
            st.markdown(f'<div class="summary-content">{transcript[:10000]}</div>', unsafe_allow_html=True)
            if len(transcript) > 10000:
                st.info("Transcript truncated for display. Full text is used for summarization.")

        with tab_chat:
            st.markdown('<div class="section-title">💬 Ask questions about the video</div>', unsafe_allow_html=True)

            # Show chat history
            for msg in st.session_state.chat_history:
                if msg["role"] == "user":
                    st.markdown(f'<div class="chat-user">🧑 {msg["content"]}</div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="chat-ai">🤖 {msg["content"]}</div>', unsafe_allow_html=True)

            # Chat input
            user_question = st.chat_input("Ask anything about this video...")
            if user_question:
                if not api_key:
                    st.error("API key not configured.")
                else:
                    st.session_state.chat_history.append({"role": "user", "content": user_question})

                    # Build conversation for API
                    system_msg = {
                        "role": "system",
                        "content": (
                            "You are a helpful assistant that answers questions about a video. "
                            "Use the following transcript to answer. If the answer is not in the "
                            "transcript, say so. Be concise and helpful.\n\n"
                            f"VIDEO TITLE: {info.get('title', 'Unknown')}\n"
                            f"TRANSCRIPT:\n{transcript[:8000]}"
                        ),
                    }
                    chat_messages = [system_msg] + [
                        {"role": m["role"], "content": m["content"]}
                        for m in st.session_state.chat_history[-6:]  # last 6 msgs for context
                    ]

                    try:
                        answer = call_euri_api(chat_messages, api_key, api_url)
                        st.session_state.chat_history.append({"role": "assistant", "content": answer})
                        st.rerun()
                    except RuntimeError as e:
                        st.error(f"❌ {e}")

    else:
        # Empty state
        st.markdown("""
        <div class="glass-card" style="text-align:center; padding: 60px 20px;">
            <div style="font-size: 3rem; margin-bottom: 16px;">🎬</div>
            <h3 style="color: #C4B5FD; margin-bottom: 8px;">Paste a YouTube URL to get started</h3>
            <p style="color: #64748B; max-width: 500px; margin: 0 auto;">
                VideoInsight AI will extract the transcript, generate an intelligent summary,
                and let you download it or chat about the content.
            </p>
        </div>
        """, unsafe_allow_html=True)

        # Feature cards
        col_a, col_b, col_c = st.columns(3)
        with col_a:
            st.markdown("""
            <div class="glass-card" style="text-align:center;">
                <div style="font-size: 2rem;">📝</div>
                <h4 style="color: #E2E8F0;">Smart Summaries</h4>
                <p style="color: #94A3B8; font-size: 0.85rem;">
                    Choose Brief, Standard, or Detailed summaries tailored to your needs
                </p>
            </div>
            """, unsafe_allow_html=True)
        with col_b:
            st.markdown("""
            <div class="glass-card" style="text-align:center;">
                <div style="font-size: 2rem;">📥</div>
                <h4 style="color: #E2E8F0;">Export Anywhere</h4>
                <p style="color: #94A3B8; font-size: 0.85rem;">
                    Download as Word document or plain text for offline reading
                </p>
            </div>
            """, unsafe_allow_html=True)
        with col_c:
            st.markdown("""
            <div class="glass-card" style="text-align:center;">
                <div style="font-size: 2rem;">💬</div>
                <h4 style="color: #E2E8F0;">Chat with Video</h4>
                <p style="color: #94A3B8; font-size: 0.85rem;">
                    Ask follow-up questions and get instant AI-powered answers
                </p>
            </div>
            """, unsafe_allow_html=True)

    # Footer
    st.markdown(
        '<div class="footer">VideoInsight AI • Powered by EURI AI & Streamlit • Open Source</div>',
        unsafe_allow_html=True,
    )


# ---------------------------------------------------------------------------
# Markdown-to-HTML helper (simplified)
# ---------------------------------------------------------------------------
def _md_to_html(md_text: str) -> str:
    """Convert markdown text to basic HTML for display."""
    import html as html_module
    text = html_module.escape(md_text)

    # Bold
    text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
    # Italic
    text = re.sub(r'\*(.*?)\*', r'<em>\1</em>', text)
    # Headers
    text = re.sub(r'^### (.*?)$', r'<h3>\1</h3>', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.*?)$', r'<h2>\1</h2>', text, flags=re.MULTILINE)
    text = re.sub(r'^# (.*?)$', r'<h1>\1</h1>', text, flags=re.MULTILINE)
    # Bullet points
    text = re.sub(r'^[-*] (.*?)$', r'<li>\1</li>', text, flags=re.MULTILINE)
    text = re.sub(r'(<li>.*?</li>\n?)+', r'<ul>\g<0></ul>', text)
    # Numbered lists
    text = re.sub(r'^\d+\. (.*?)$', r'<li>\1</li>', text, flags=re.MULTILINE)
    # Line breaks
    text = text.replace('\n\n', '</p><p>')
    text = text.replace('\n', '<br>')
    text = f'<p>{text}</p>'
    return text


# ---------------------------------------------------------------------------
if __name__ == "__main__":
    main()
